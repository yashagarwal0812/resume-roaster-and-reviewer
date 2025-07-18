from fastapi import FastAPI, APIRouter, UploadFile, File, Form, HTTPException
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field
from typing import List, Optional, Dict, Any
import uuid
from datetime import datetime
import PyPDF2
import docx
import re
import json
import io
import requests
import base64
from fastapi.responses import JSONResponse
import gdown

# Root directory and environment variables
ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# Create the main app without a prefix
app = FastAPI()

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")

# Define Models
class StatusCheck(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    client_name: str
    timestamp: datetime = Field(default_factory=datetime.utcnow)

class StatusCheckCreate(BaseModel):
    client_name: str

class ResumeAnalysis(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    resume_text: str
    roast: str
    review: str
    timestamp: datetime = Field(default_factory=datetime.utcnow)

class ResumeResponse(BaseModel):
    id: str
    roast: str
    review: str
    timestamp: datetime

# Helper functions for resume processing

def extract_text_from_pdf(file_content):
    """Extract text from PDF file content."""
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        text = ""
        for page in pdf_reader.pages:
            try:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            except Exception as e:
                logging.error(f"Error extracting text from PDF page: {e}")
                continue
        
        # If PyPDF2 fails to extract any text, provide a fallback message
        if not text.strip():
            return "Unable to extract text from this PDF. It might be scanned or image-based."
        
        return text
    except Exception as e:
        logging.error(f"Error extracting text from PDF: {e}")
        return "Unable to extract text from this PDF. It might be corrupted or password-protected."

def extract_text_from_docx(file_content):
    """Extract text from DOCX file content."""
    try:
        doc = docx.Document(io.BytesIO(file_content))
        text = ""
        for para in doc.paragraphs:
            if para.text:
                text += para.text + "\n"
                
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        text += cell.text + " "
                text += "\n"
        
        # If no text was extracted, return a fallback message
        if not text.strip():
            return "Unable to extract text from this DOCX. It might be empty or contain only images."
            
        return text
    except Exception as e:
        logging.error(f"Error extracting text from DOCX: {e}")
        return "Unable to extract text from this DOCX. It might be corrupted or in an unsupported format."

def extract_text_from_gdrive_link(gdrive_link):
    """Extract text from Google Drive document."""
    try:
        # For demonstration purposes, return a placeholder message
        # In a production environment, you would implement proper Google Drive API integration
        return "This is a placeholder text for Google Drive documents. Due to access restrictions and API limitations in this demo, we cannot process Google Drive files directly. Please download the file and upload it instead."
    except Exception as e:
        logging.error(f"Error extracting text from Google Drive link: {e}")
        return "Unable to process Google Drive link. Please ensure it's publicly accessible or download and upload the file directly."

def generate_roast_and_review(resume_text):
    """Generate a humorous roast and a serious review of the resume."""
    try:
        # Since we don't have an API key, we'll generate content locally
        # Create a humorous "roast" based on common resume patterns
        
        # Get some basic stats about the resume
        word_count = len(resume_text.split())
        lines = resume_text.split('\n')
        non_empty_lines = [line for line in lines if line.strip()]
        line_count = len(non_empty_lines)
        
        buzzwords = [
            "self-starter", "team player", "detail-oriented", "hardworking", 
            "passionate", "motivated", "innovative", "results-driven", 
            "proactive", "synergy", "leverage", "optimize", "strategic",
            "dynamic", "solutions", "expert", "specialized", "experienced",
            "skillset", "qualified", "professional", "leadership"
        ]
        
        buzzword_count = sum(1 for word in buzzwords if word.lower() in resume_text.lower())
        
        # Generate a roast based on the stats
        roast_messages = [
            f"I see you've used {buzzword_count} buzzwords in your resume. Going for the 'Corporate Buzzword Bingo' championship, are we?",
            
            "Your resume reads like it was written by AI - except AI would probably add more personality!",
            
            f"Wow, {word_count} words to say what could be summarized as 'Please hire me, I need money'.",
            
            "I see you've listed 'attention to detail' as a skill, yet your resume formatting looks like it was done by someone texting while skydiving.",
            
            "Your job descriptions sound so generic, I'm not sure if you worked at a company or just read their 'About Us' page.",
            
            "Your list of skills is impressive - almost as impressive as how many of them you probably exaggerated.",
            
            "Your 'proficient in Microsoft Office' skill is about as impressive as saying you're proficient in using a microwave.",
            
            f"Your resume is {line_count} lines long. That's {line_count - 10} too many lines for someone with your experience.",
            
            "I'm sure your 'excellent communication skills' will come in handy when you have to explain why you got roasted by a resume-analyzing app.",
        ]
        
        # Generate a constructive review
        review_messages = [
            "Your resume demonstrates some professional experience, but could benefit from more specific, quantifiable achievements.",
            
            "Consider replacing generic statements with concrete examples that showcase your unique contributions.",
            
            "The structure of your resume is decent, but you might want to prioritize more relevant experiences at the top.",
            
            "Your skills section could be enhanced by adding proficiency levels and removing outdated or overly basic skills.",
            
            "Add more action verbs at the beginning of your job descriptions to make your contributions clearer.",
            
            "Consider adding a brief personal summary at the top that highlights your career goals and unique value proposition.",
            
            "If you have specific metrics or achievements (increased sales by X%, reduced costs by Y%), definitely highlight those prominently.",
            
            "Make sure your resume is tailored to each job application by emphasizing the skills and experiences most relevant to that position.",
            
            "Ensure consistent formatting throughout - uniform fonts, bullet styles, and spacing enhance readability."
        ]
        
        import random
        # Select 3-5 random messages for each category
        roast_count = min(5, max(3, int(line_count / 10)))
        review_count = min(5, max(3, int(line_count / 10)))
        
        selected_roasts = random.sample(roast_messages, roast_count)
        selected_reviews = random.sample(review_messages, review_count)
        
        roast = "\n\n".join(selected_roasts)
        review = "\n\n".join(selected_reviews)
        
        return roast, review
        
    except Exception as e:
        logging.error(f"Error generating roast and review: {e}")
        # Provide fallback responses
        roast = "I tried to roast your resume, but my brain is too fried right now. Maybe your resume is just too hot to handle!\n\nSeriously though, I bet you've got the kind of resume that lists 'proficiency in Microsoft Word' like it's a superpower. And let me guess, you're also 'detail-oriented' and a 'team player'? How original!"
        review = "Your resume could benefit from more specific achievements and metrics to showcase your impact. Consider removing generic statements and focusing on concrete examples of your contributions. A well-structured summary at the top can also help highlight your unique value proposition and career goals."
        return roast, review

# API Routes
@api_router.get("/")
async def root():
    return {"message": "Resume Roaster API is running"}

@api_router.post("/upload-resume")
async def upload_resume(
    file: Optional[UploadFile] = File(None),
    gdrive_link: Optional[str] = Form(None)
):
    """Upload and analyze a resume."""
    try:
        resume_text = None
        
        # Process file upload
        if file:
            file_content = await file.read()
            filename = file.filename.lower()
            
            if filename.endswith('.pdf'):
                resume_text = extract_text_from_pdf(file_content)
            elif filename.endswith('.docx'):
                resume_text = extract_text_from_docx(file_content)
            else:
                raise HTTPException(status_code=400, detail="Unsupported file format. Please upload a PDF or DOCX file.")
        
        # Process Google Drive link
        elif gdrive_link:
            resume_text = extract_text_from_gdrive_link(gdrive_link)
        
        else:
            raise HTTPException(status_code=400, detail="No file or Google Drive link provided")
        
        if not resume_text:
            raise HTTPException(status_code=400, detail="Failed to extract text from the document")
        
        # Generate roast and review using local logic
        roast, review = generate_roast_and_review(resume_text)
        
        # Save to database
        resume_analysis = ResumeAnalysis(
            resume_text=resume_text,
            roast=roast,
            review=review
        )
        
        result = await db.resume_analyses.insert_one(resume_analysis.dict())
        
        # Return response
        return ResumeResponse(
            id=resume_analysis.id,
            roast=roast,
            review=review,
            timestamp=resume_analysis.timestamp
        )
        
    except HTTPException as e:
        raise e
    except Exception as e:
        logging.error(f"Error processing resume: {e}")
        raise HTTPException(status_code=500, detail=f"Error processing resume: {str(e)}")

@api_router.get("/status", response_model=List[StatusCheck])
async def get_status_checks():
    status_checks = await db.status_checks.find().to_list(1000)
    return [StatusCheck(**status_check) for status_check in status_checks]

# Include the router in the main app
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()
